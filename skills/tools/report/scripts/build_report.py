#!/usr/bin/env python3
# /// script
# requires-python = ">=3.9"
# dependencies = ["python-docx"]
# ///
"""Assemble a near-submission .docx report from a JSON manifest.

Usage: uv run build_report.py manifest.json [-o report.docx]

Run via `uv run` so python-docx resolves into an isolated, cached env. The
agent writes the manifest from its result package; this script only assembles
it deterministically (headings, paragraphs, tables, embedded figures, numbered
captions). It does NOT invent content.

Manifest schema (JSON):
{
  "title": "Computational response to peer review",
  "subtitle": "optional line under the title",
  "sections": [
    {
      "heading": "R2.8 - interfacial Pt valence state",
      "level": 1,                       # 1..3
      "paragraphs": ["plain text ...", "..."],
      "tables": [
        {"caption": "Relative energies (eV)",
         "columns": ["Model", "E_ads (eV)", "Bader q(Pt)"],
         "rows": [["Pt4 cluster", "-1.97", "+0.05"], ...]}
      ],
      "figures": [
        {"path": "figs/pt4_top_side_panel.png",
         "caption": "(a) Top view (ortho); (b) zoomed side view (ortho).",
         "width_in": 6.2}
      ]
    }
  ]
}

Tables get a caption ABOVE; figures a caption BELOW (journal convention).
Captions are auto-numbered (Table 1.., Figure 1..).

Advisory check (does not fail the build): a numeric table cell with |value| > 50
is flagged on stderr as a likely *total* energy — humans read relative energies
(adsorption/binding energies, barriers), not raw TOTEN. See
`knowledge/scientific-visualization.md` and `references/validation.md`.
"""

import argparse
import json
import os
import re
import sys

TOTAL_ENERGY_GUESS = (
    50.0  # |eV| above which an energy-column value smells like a total energy
)
# only flag columns that look like an energy (avoids frequencies in cm^-1, temperatures, etc.)
_ENERGY_HEADER = re.compile(
    r"toten|total\s*energy|free\s*energy|\benergy\b|\beV\b|\be[_ ]?tot\b|\bE\s*\(",
    re.IGNORECASE,
)


def looks_like_total_energy(cell, header=""):
    if not _ENERGY_HEADER.search(str(header)):
        return False
    try:
        return abs(float(str(cell).replace(",", ""))) > TOTAL_ENERGY_GUESS
    except (TypeError, ValueError):
        return False


def main():
    ap = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    ap.add_argument("manifest", help="path to manifest.json")
    ap.add_argument("-o", "--out", default="report.docx")
    args = ap.parse_args()

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        sys.exit(
            "python-docx is required: run via `uv run build_report.py ...` "
            "(or pip install python-docx)"
        )

    with open(args.manifest) as fh:
        m = json.load(fh)

    base = os.path.dirname(os.path.abspath(args.manifest))
    doc = Document()
    doc.add_heading(m.get("title", "Report"), level=0)
    if m.get("subtitle"):
        sub = doc.add_paragraph(m["subtitle"])
        sub.runs[0].italic = True

    tbl_n = fig_n = 0
    warnings = []

    def caption(text, label, n):
        p = doc.add_paragraph()
        r = p.add_run(f"{label} {n}. {text}")
        r.italic = True
        r.font.size = Pt(9)
        return p

    for s in m.get("sections", []):
        doc.add_heading(s.get("heading", ""), level=min(max(s.get("level", 1), 1), 3))
        for para in s.get("paragraphs", []):
            doc.add_paragraph(para)

        for t in s.get("tables", []):
            tbl_n += 1
            if t.get("caption"):
                caption(t["caption"], "Table", tbl_n)
            cols = t.get("columns", [])
            rows = t.get("rows", [])
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            for j, c in enumerate(cols):
                table.rows[0].cells[j].text = str(c)
            for ri, row in enumerate(rows):
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
                    if looks_like_total_energy(val, cols[j] if j < len(cols) else ""):
                        warnings.append(
                            f"Table {tbl_n} row {ri + 1} col '{cols[j] if j < len(cols) else j}'"
                            f" = {val}: looks like a TOTAL energy — report a relative energy"
                            " (E_ads / E_bind / barrier / dG) instead."
                        )

        for f in s.get("figures", []):
            path = f.get("path", "")
            full = path if os.path.isabs(path) else os.path.join(base, path)
            fig_n += 1
            if not os.path.isfile(full):
                doc.add_paragraph(f"[missing figure: {path}]")
                warnings.append(f"Figure {fig_n}: file not found: {full}")
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(full, width=Inches(f.get("width_in", 5.5)))
            if f.get("caption"):
                caption(f["caption"], "Figure", fig_n)

    doc.save(args.out)
    print(
        f"wrote {args.out}: {len(m.get('sections', []))} sections, "
        f"{tbl_n} tables, {fig_n} figures"
    )
    for w in warnings:
        print(f"ADVISORY: {w}", file=sys.stderr)
    # advisory only — never fail the build on the heuristic
    sys.exit(0)


if __name__ == "__main__":
    main()
